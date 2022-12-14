# -*- coding: utf-8 -*-

"""
- OK --> Detallar en el gráfico cuál es el punto de inicio y punto de fin de un viaje
- OK --> Configurar el zoom automático.
- Calcular también lugares cercanos para final del trayecto
- Poner etiquetas de lugares frecuentes: ej: Marriott, Oficina, York, aeropuertos, etc.
- Poner el lugar cercano en el word
- Correo con notificación de ejecución
- OK --> Informe de sitios turísticos.
- OK --> Agregar imágenes más pequeñas en zooms lejanos
- OK --> Configurar cuando el reporte se descargue en español.
- OK --> Tomar la fecha desde el nombre del reporte, no la del día anterior. Actualizarla en el word.
"""

import os
import config
import logging
import pandas as pd
import sqldf

import plotly.graph_objects as go

from io import BytesIO
from docx import Document
from docx.shared import Inches
from haversine import haversine, haversine_vector

def read_data(report_path: str = config.report_path):
    global df
    
    df = pd.read_excel(report_path, header=3)
    df.rename(mapper=config.spa_eng_cols, inplace=True, axis=1)
    
    return df

def create_vehicle_type_folders(vehicle_types = config.vehicle_types):
    vehicle_types.append('Otros')
    for vtype in vehicle_types:
        vt_path = os.path.join(config.control_path, vtype)
        if not os.path.isdir(vt_path):
            os.mkdir(vt_path)

def create_logger(control_path = config.control_path):
    log_path = os.path.join(control_path, '0_Logs')
    if not os.path.isdir(log_path):
        os.mkdir(log_path)
        
    logging.basicConfig(filename=f'{log_path}\{config.start_date}-{config.end_date}.log', encoding='utf-8',
                        format='%(asctime)s: %(message)s', datefmt='%Y.%m.%d %H:%M %p',
                        level=logging.DEBUG)

def clean_data(df):
    filter_1 = df['Vehicle plate number'].notna()
    filter_2 = df['#'] != '#'
    df = df[filter_1 & filter_2]
    
    df['Trip State'] = df['Trip State'].replace({'Conducir': 'Driving', 'Estacionamiento': 'Parking'})
    df['Start time'] = pd.to_datetime(df['Start time'])
    df['End time'] = pd.to_datetime(df['End time'])
    df['Duration_mins'] = (df['End time'] - df['Start time'])
    df['Duration_mins'] = df['Duration_mins'].apply(lambda x: x.total_seconds() / 60)
    df['Mileage (KM)'] = df['Mileage (KM)'].replace('-', 0).astype(float)
    df['Start location'] = df['Start location'].str.replace('N', '').str.replace('W', '')
    df['End location'] = df['End location'].str.replace('N', '').str.replace('W', '')
    df[['Latitud_Inicio', 'Longitud_Inicio']] = df['Start location'].str.split(',', expand=True)
    df[['Latitud_Fin', 'Longitud_Fin']] = df['End location'].str.split(',', expand=True)
    location_cols = ['Latitud_Inicio', 'Longitud_Inicio', 'Latitud_Fin', 'Longitud_Fin']
    df[location_cols] = df[location_cols].fillna(0.0).replace('-', 0.0).astype(float)
    df[['Longitud_Inicio', 'Longitud_Fin']] = df[['Longitud_Inicio', 'Longitud_Fin']].mul(-1)
    df = df.reset_index(drop=True)
    df.drop(['#', 'Duration', 'Start location', 'End location'], axis=1, inplace=True)
    
    logging.info('Datos limpiados con éxito.')
    
    return df

def common_places(df, df_locations = config.df_locations):
    df['Ubicacion_Inicio'] = list(zip(df['Latitud_Inicio'], df['Longitud_Inicio']))
    df_locations['Acopio'] = df_locations['Acopio'].fillna('NO')
    df_locations['Ubicacion_Inicio'] = list(zip(df_locations['Latitud'], df_locations['Longitud']))
    array_1 = df_locations['Ubicacion_Inicio'].tolist()
    array_2 = df['Ubicacion_Inicio'].tolist()
    
    dists = haversine_vector(array_1, array_2, unit='km', comb=True)
    df_dists = pd.DataFrame(dists, columns=df_locations['Location'])
    df['Lugar_Inicio'] = df_dists.idxmin(axis=1)
    df['Dist'] = df_dists.min(axis=1)
    df = df.merge(df_locations[['Location', 'Perimetro KM', 'Acopio']],
                  left_on=['Lugar_Inicio'], right_on=['Location'], how='left')
    df['Lugar_Inicio'] = df.apply(lambda x: x['Lugar_Inicio'] if x['Perimetro KM'] > x['Dist'] else None, axis=1)
    
    df.drop(['Ubicacion_Inicio', 'Location', 'Perimetro KM'], axis=1, inplace=True)
    
    logging.info('Lugares cercanos calculados con éxito.')

    return df

def identificar_acopios(df, df_vehicles = config.df_vehicles):
    
    df = df.merge(df_vehicles, left_on='Vehicle plate number', right_on='Placa', how='left')
    df['Tipo'].fillna('Otros', inplace=True)
    
    bool_i = df['Lugar_Inicio'].notna()
    bool_ii = df['Acopio']=='SI'
    bool_iii = df['Trip State']=='Parking'
    df_acopios = df[bool_i & bool_ii & bool_iii]
    df_acopios = df_acopios.groupby(['Vehicle plate number', 'Tipo', 'Lugar_Inicio'])['Duration_mins'].sum().reset_index()
    df_acopios.sort_values(by=['Duration_mins'], ascending=False, inplace=True)
    
    logging.info('Acopios identificados correctamente.')
    
    return df, df_acopios
    
def identificar_descansos(df, df_schedules = config.df_schedules):
    df_schedules = df_schedules.melt(id_vars=['Conductor'], var_name='FECHA', value_name='HORARIO')
    df_schedules = df_schedules[df_schedules['HORARIO']=='DESCANSA']
    df_schedules['Y'] = df_schedules['FECHA'].dt.year
    df_schedules['M'] = df_schedules['FECHA'].dt.month
    df_schedules['D'] = df_schedules['FECHA'].dt.day
    st = df['Start time']
    df['Y'], df['M'], df['D'] = st.dt.year, st.dt.month, st.dt.day
    df = df.merge(df_schedules, on=['Conductor', 'Y', 'M', 'D'], how='left')
    
    logging.info('Descansos identificados correctamente.')
    
    return df

def aggregate_metrics(df):
    global df_drive_agg, df_park_agg, df_unique
    
    df_drive_agg = sqldf.run(config.driving_agg)
    df_park_agg = sqldf.run(config.parking_agg)
    df_agg = sqldf.run(config.join_agg)
    df_agg.drop(['index'], axis=1, inplace=True)
    
    df_unique = df[['Vehicle plate number', 'Lugar_Inicio']].drop_duplicates(['Vehicle plate number', 'Lugar_Inicio'])
    df_sql = sqldf.run(config.query_places)
    
    df_agg = df_agg.merge(df_sql, on=['Placa'], how='left')
    logging.info('Métricas por placa calculadas con éxito.')
    
    return df_agg
    
def plot_heatmap_trips(df):
    fig = go.Figure(go.Densitymapbox(lat=df['Latitud_Inicio'], lon=df['Longitud_Inicio'], radius=5))
    
    lat_center = (df['Latitud_Inicio'].max() + df['Latitud_Inicio'].min())/2
    lon_center = (df['Longitud_Inicio'].max() + df['Longitud_Inicio'].min())/2
    fig.update_layout(margin = {'b': 0, 'l': 0, 'r': 0, 't':0},
                      mapbox = {'center': {'lat': lat_center, 'lon': lon_center},
                                'style': "open-street-map", 'zoom': 10})
    
    fig.write_html(r".\trip_metrics\trips_mapa_calor.html")
    logging.info('Mapa de calor creado con éxito.')

def plot_all_trips(df_plate):
    longs = df_plate['Longitud_Inicio'].tolist()
    longs.append(df_plate['Longitud_Fin'].tolist()[-1])
    lats = df_plate['Latitud_Inicio'].tolist()
    lats.append(df_plate['Latitud_Fin'].tolist()[-1])
    fig = go.Figure(go.Scattermapbox(mode = "markers+lines", marker = {'size': 9},
                                    lon = longs, lat = lats))
    min_min = (df_plate['Latitud_Inicio'].min(), df_plate['Longitud_Inicio'].min())
    max_max = (df_plate['Latitud_Inicio'].max(), df_plate['Longitud_Inicio'].max())
    zoom = set_plot_zoom(min_min, max_max)
    
    lat_center = (df_plate['Latitud_Inicio'].max() + df_plate['Latitud_Inicio'].min())/2
    lon_center = (df_plate['Longitud_Inicio'].max() + df_plate['Longitud_Inicio'].min())/2
    fig.update_layout(margin = {'b': 0, 'l': 0, 'r': 0, 't':0},
                      mapbox = {'center': {'lat': lat_center, 'lon': lon_center},
                                'style': "open-street-map", 'zoom': zoom})
    fig_png = BytesIO(fig.to_image(format="png"))
    
    document.add_picture(fig_png, width=Inches(6.0))

def plot_single_trips(df_plate):
    trip_num = 0
    for index, row in df_plate.iterrows():
        trip_num += 1
        document.add_page_break()
        p = document.add_paragraph()
        p.add_run(f'Viaje {trip_num}').bold = True
        document.add_paragraph(f'{row["Start time"]} - {row["End time"]}')
        document.add_paragraph('Duración: %.1f minutos.' %row["Duration_mins"])
        document.add_paragraph('Total KM: %.1f' %row["Mileage (KM)"])
        fig = go.Figure(go.Scattermapbox(mode = "markers+lines",
                                         marker = {'size': [12,12], 'opacity': 1,
                                                   'color': ['blue', 'red']},
                                         line = {'color': 'black'},
                                         lon = [row['Longitud_Inicio'], row['Longitud_Fin']],
                                         lat = [row['Latitud_Inicio'], row['Latitud_Fin']]))
        lat_center = (row['Latitud_Inicio'] + row['Latitud_Fin'])/2
        lon_center = (row['Longitud_Inicio'] + row['Longitud_Fin'])/2
        min_min = (min([row['Latitud_Inicio'], row['Latitud_Fin']]), min([row['Longitud_Inicio'], row['Longitud_Fin']]))
        max_max = (max([row['Latitud_Inicio'], row['Latitud_Fin']]), max([row['Longitud_Inicio'], row['Longitud_Fin']]))
        zoom = set_plot_zoom(min_min, max_max)
        fig.update_layout(margin = {'b': 0, 'l': 0, 'r': 0, 't':0},
                          mapbox = {'center': {'lat': lat_center,'lon': lon_center},
                                    'style': "open-street-map", 'zoom': zoom})
        fig_png = BytesIO(fig.to_image(format="png"))
        document.add_picture(fig_png, width=Inches(6.0))
        
        if zoom <= 12:
            plot_start_end_location([row['Latitud_Inicio'], row['Latitud_Fin']],
                                    [row['Longitud_Inicio'], row['Longitud_Fin']])

def plot_start_end_location(lat: list, lon: list):
    
    p = document.add_paragraph()
    for i, ll in enumerate(tuple(zip(lat, lon))):
        color = 'blue'
        if i == 1:
            color = 'red'
        
        fig = go.Figure()
        fig.add_trace(go.Scattermapbox(mode='markers' , lat = [ll[0]], lon = [ll[1]], 
                                         marker = {'size': 18, 'opacity': 1, 'color': color}))
        fig.update_layout(margin = {'b': 0, 'l': 0, 'r': 0, 't':0},
                          mapbox = {'center': {'lat': ll[0],'lon': ll[1]},
                                    'style': "open-street-map", 'zoom': 16})
        fig_png = BytesIO(fig.to_image(format="png"))
        p.add_run().add_picture(fig_png, width=Inches(3.0))
    

def set_plot_zoom(min_min, max_max):
    dist = haversine(min_min, max_max)
    zoom = 10
    for i in config.zoom_list:
        if dist < i[0]:
            zoom = i[1]
            break
    #document.add_paragraph(f'TEMPORAL: Distancia: {dist} ZOOM: {zoom}')
    return zoom

def create_trips_docx(df):
    global document
    
    for plate in df['Vehicle plate number'].unique():
        document = Document()
        document.add_picture(r".\parametros_control\img\img_control.png", width=Inches(1.0))
        p = document.add_paragraph()
        p.add_run(f'INFORME VERIFICACIÓN DE GPS VERSIÓN 01 - {config.today}').bold = True
        document.add_paragraph().add_run(f'PLACA: {plate}')
        document.add_paragraph().add_run(f'FECHA: {config.start_date}')
        
        vtype = df[df['Vehicle plate number']==plate]['Tipo'].values[0]
        if vtype not in config.vehicle_types:
            vtype = 'Otros'
            
        plate_path = os.path.join(config.control_path, vtype, f'{plate}')
        if not os.path.isdir(plate_path):
            os.mkdir(plate_path)
        
        df_plate = df[(df['Vehicle plate number']==plate) & (df['Trip State']=='Driving')]
        
        if df_plate.shape[0] > 0:
            plot_all_trips(df_plate)
            plot_single_trips(df_plate)  
            document.save(os.path.join(plate_path, f'{config.start_date}-{config.end_date}_{plate}.docx'))
        
        print(f'PROCESADA: {plate} - {vtype}')
        logging.info(f'PROCESADA: {plate} - {vtype}')

def export_results(dfs: list, sheet_names: list):
    writer = pd.ExcelWriter(rf".\trip_metrics\{config.start_date}-{config.end_date}_trip_metrics.xlsx", engine='openpyxl')
    for df, name in zip(dfs, sheet_names):
        df.to_excel(writer, sheet_name=name, index=False)
        writer.save()
    logging.info('Resultados exportados correctamente.')

def main_gps():
    create_logger()
    logging.info('INICIO EJECUCIÓN')
    df = read_data()
    create_vehicle_type_folders(config.vehicle_types)
    df = clean_data(df)
    df = common_places(df)
    df, df_acopios = identificar_acopios(df)
    df = identificar_descansos(df)
    df_agg = aggregate_metrics(df)
    plot_heatmap_trips(df)
    create_trips_docx(df)
    export_results([df, df_agg, df_acopios],
                   ['DATOS', 'MÉTRICAS', 'ACOPIOS'])
    logging.info('FIN EJECUCIÓN')
    
if __name__ == '__main__':
    main_gps()
